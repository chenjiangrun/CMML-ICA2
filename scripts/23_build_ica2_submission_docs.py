from __future__ import annotations

import csv
import re
from pathlib import Path

from PIL import Image, ImageChops, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
FIG_DIR = ROOT / "figures" / "real"
COMPOSITE_FIG = FIG_DIR / "main_report_composite_figure.png"


TITLE = "Benchmarking Doublet Detection Methods in Single-cell RNA-seq Data"
STUDENT = "3127"
COURSE = "CMML3 ICA2 Miniproject 7"


ABSTRACT = (
    "Doublets are a practical problem in single-cell RNA-seq because one barcode "
    "can contain RNA from more than one cell. I compared Scrublet, "
    "DoubletFinder, and scDblFinder on 16 labelled real datasets and a small "
    "set of simulations. DoubletFinder gave the best hard-call F1, while "
    "scDblFinder ranked likely doublets better. The results mainly show that "
    "threshold choice matters."
)


MAIN_SECTIONS = [
    (
        "Background and aim",
        (
            "In droplet-based single-cell RNA-seq, two cells can occasionally be "
            "captured together and counted as one barcode. These doublets can "
            "shift clustering, marker interpretation, and differential "
            "expression. Computational methods try to detect them from unusual "
            "expression profiles, but doublets made from similar cell types are "
            "still hard to recognise. I compared three commonly used methods, "
            "Scrublet "
            "(Wolock, Lopez and Klein, 2019), DoubletFinder (McGinnis, Murrow "
            "and Gartner, 2019), and scDblFinder (Germain et al., 2022), under "
            "the same processing and evaluation workflow."
        ),
    ),
    (
        "Methods",
        (
            "I used all 16 real datasets from the Xi and Li benchmark archive, "
            "which provide ground-truth doublet labels (Xi and Li, 2020; Xi and "
            "Li, 2021). To keep the analysis runnable on a local machine, large "
            "datasets were sampled to 3000 cells, with doublets preserved up to "
            "a 35% cap. The methods were run on raw count matrices. For the "
            "downstream checks I used filtering, normalisation, highly variable "
            "genes, PCA, Leiden clustering, and UMAP. I measured precision, "
            "recall, F1, AUROC, AUPRC, and runtime. Because thresholds differ "
            "between methods, I also ranked cells by score and called the top N "
            "cells, where N was the true processed doublet count. Simulations "
            "were used only as supporting checks for doublet rate, marker "
            "separation, and homotypic versus heterotypic doublets "
            "(Supplementary Figures 1-3)."
        ),
    ),
    (
        "Results",
        (
            "The real-data results were mixed rather than giving a single clear "
            "winner. DoubletFinder had the highest mean thresholded F1 (0.472) "
            "and recall (0.423). scDblFinder had the highest mean AUROC "
            "(0.780), AUPRC (0.634), and precision (0.777). Scrublet was the "
            "fastest method, but its recall varied substantially between datasets. "
            "DoubletFinder was best by F1 on 8 of 16 real datasets, scDblFinder "
            "on 4, and Scrublet on 4 (Figure 1A). Micro-averaging gave the same "
            "hard-call pattern: DoubletFinder F1=0.509, Scrublet F1=0.484, and "
            "scDblFinder F1=0.333. The score-based comparison told a different "
            "story (Figure 1B-C). When each method was allowed to call the same "
            "number of doublets, scDblFinder had the highest mean calibrated F1 "
            "(0.602), followed by Scrublet (0.534) and DoubletFinder (0.498). "
            "This suggests that scDblFinder often ranked true doublets well but "
            "called too few cells by default; false negatives are summarised in "
            "Supplementary Figure 4."
        ),
    ),
    (
        "Downstream impact",
        (
            "I used pbmc-1A-dm as a representative downstream example. After "
            "removing scDblFinder-predicted doublets, the UMAP layout changed "
            "and marker-effect estimates shifted. The mean top-10 marker "
            "overlap was still high (0.920), so the main marker identities were "
            "not lost, but shared markers had a mean absolute log-fold-change "
            "shift of 0.524. In this dataset, doublet removal therefore changed "
            "effect sizes more than marker identities (Figure 1D). The "
            "simulation results also showed why some doublets were missed: "
            "heterotypic doublets were easier to detect than homotypic doublets. "
            "Marker stability is shown in Supplementary Figure 5."
        ),
    ),
    (
        "Discussion",
        (
            "The main point I take from these results is that score ranking and "
            "hard filtering should not be treated as the same task. scDblFinder "
            "looked useful for ranking suspicious barcodes, while DoubletFinder "
            "gave stronger thresholded calls in several real datasets. Doublets "
            "also cannot be removed completely by experimental design: lowering "
            "cell loading can reduce them, but it also reduces throughput. I "
            "would not automatically ignore a low doublet rate if rare cell "
            "states or differential expression are important. The main "
            "limitations are the 3000-cell subsampling, the use of benchmark "
            "doublet rates for hard-call settings, and the local "
            "Seurat-v5-compatible DoubletFinder implementation. In practice, I "
            "would rank cells with scDblFinder, compare removal decisions with "
            "DoubletFinder, and inspect rare populations before deleting them."
        ),
    ),
]


def create_composite_figure() -> Path:
    panels = [
        ("A", FIG_DIR / "f1_by_dataset_heatmap.png"),
        ("B", FIG_DIR / "auprc_auroc_comparison.png"),
        ("C", FIG_DIR / "calibrated_f1_precision_recall_comparison.png"),
        ("D", FIG_DIR / "umap_before_after_doublet_removal.png"),
    ]
    cell_w, cell_h = 1050, 760
    margin = 40
    label_h = 44
    sheet = Image.new("RGB", (cell_w * 2 + margin * 3, cell_h * 2 + margin * 3), "white")
    draw = ImageDraw.Draw(sheet)
    try:
        font = ImageFont.truetype("arial.ttf", 34)
    except OSError:
        font = ImageFont.load_default()

    def trim_whitespace(image: Image.Image) -> Image.Image:
        bg = Image.new("RGB", image.size, "white")
        diff = Image.eval(ImageChops.difference(image, bg), lambda px: 255 if px > 12 else 0)
        bbox = diff.getbbox()
        if not bbox:
            return image
        left, top, right, bottom = bbox
        pad = 18
        return image.crop(
            (
                max(0, left - pad),
                max(0, top - pad),
                min(image.width, right + pad),
                min(image.height, bottom + pad),
            )
        )

    for idx, (label, path) in enumerate(panels):
        row, col = divmod(idx, 2)
        x = margin + col * (cell_w + margin)
        y = margin + row * (cell_h + margin)
        draw.text((x, y), label, fill=(31, 78, 121), font=font)
        image = trim_whitespace(Image.open(path).convert("RGB"))
        max_w = cell_w
        max_h = cell_h - label_h
        image.thumbnail((max_w, max_h), Image.Resampling.LANCZOS)
        px = x + (max_w - image.width) // 2
        py = y + label_h + (max_h - image.height) // 2
        sheet.paste(image, (px, py))

    COMPOSITE_FIG.parent.mkdir(parents=True, exist_ok=True)
    sheet.save(COMPOSITE_FIG, quality=95)
    return COMPOSITE_FIG


REFERENCES = [
    "Germain, P.L., Lun, A., Garcia Meixide, C., Macnair, W. and Robinson, M.D. "
    "(2022) 'Doublet identification in single-cell sequencing data using "
    "scDblFinder', F1000Research, 10, p. 979. doi: 10.12688/f1000research.73600.2.",
    "McGinnis, C.S., Murrow, L.M. and Gartner, Z.J. (2019) 'DoubletFinder: "
    "doublet detection in single-cell RNA sequencing data using artificial "
    "nearest neighbors', Cell Systems, 8(4), pp. 329-337.e4. doi: "
    "10.1016/j.cels.2019.03.003.",
    "Wolock, S.L., Lopez, R. and Klein, A.M. (2019) 'Scrublet: computational "
    "identification of cell doublets in single-cell transcriptomic data', Cell "
    "Systems, 8(4), pp. 281-291.e9. doi: 10.1016/j.cels.2018.11.005.",
    "Xi, N.M. and Li, J.J. (2020) Benchmarking computational doublet-detection "
    "methods for single-cell RNA sequencing data. Zenodo. Available at: "
    "https://doi.org/10.5281/zenodo.4562782.",
    "Xi, N.M. and Li, J.J. (2021) 'Benchmarking computational doublet-detection "
    "methods for single-cell RNA sequencing data', Cell Systems, 12(2), pp. "
    "176-194.e6. doi: 10.1016/j.cels.2020.11.008.",
]


SUPP_METHODS = (
    "Real benchmark data were downloaded from the Xi and Li Zenodo archive and "
    "verified against the official MD5 checksum. The archive contained 16 RDS "
    "datasets with count matrices and ground-truth doublet labels. Each dataset "
    "was converted into a common Matrix Market format with cell and gene "
    "metadata. Datasets larger than 3000 cells were sampled with seed 7, "
    "preserving all positives where possible and capping processed doublet "
    "fraction at 35% to avoid unrealistic positive dominance. Detector inputs "
    "used raw count matrices. Scrublet was run in Python with the processed "
    "ground-truth doublet rate as the expected rate. scDblFinder was run in R "
    "with package defaults. DoubletFinder was run through a Seurat-v5-compatible "
    "pANN workflow because the installed upstream package had a metadata "
    "compatibility issue with local Seurat; the workflow retained artificial "
    "doublet generation, joint PCA, nearest-neighbour pANN scoring, "
    "expected-rate thresholding, and homotypic adjustment. Metrics were computed "
    "with doublets as the positive class. Precision, recall, and F1 used each "
    "method's hard calls; AUROC and AUPRC used continuous scores. A calibrated "
    "top-N comparison ranked cells by score and called the true number of "
    "processed doublets per dataset. Downstream analysis used Scanpy "
    "normalisation, highly variable genes, PCA, neighbours, Leiden clustering, "
    "and UMAP. Controlled simulations generated known singlet, heterotypic, and "
    "homotypic labels to test rate and composition effects. The full code, "
    "parameters, quality-control checks, and reproducibility manifest are "
    "included in the submitted repository package."
)


REFLECTION = (
    "This project was more difficult than I expected because the software setup "
    "and data download took almost as much effort as the analysis. I had to "
    "check file integrity, keep the assignment brief, and later rebuild the "
    "report when I found the strict ICA2 length rules. The scientific point I "
    "will remember is that a good doublet score is not the same as a good "
    "default threshold. If I extended the work, I would run the full datasets on "
    "a server, choose thresholds without using true labels, and repeat the "
    "marker-impact analysis across more real datasets."
)


SUPP_FIGURES = [
    (
        "Supplementary Figure 1",
        FIG_DIR.parent / "pipeline_workflow.png",
        "Reproducible workflow from controlled simulation and real-data "
        "preprocessing to method execution, evaluation, figures, and reporting.",
    ),
    (
        "Supplementary Figure 2",
        FIG_DIR.parent / "f1_precision_recall_comparison.png",
        "Precision, recall, and F1 across the controlled simulation benchmark.",
    ),
    (
        "Supplementary Figure 3",
        FIG_DIR.parent / "doublet_type_recall_comparison.png",
        "Detection of homotypic versus heterotypic doublets in controlled "
        "simulation scenarios.",
    ),
    (
        "Supplementary Figure 4",
        FIG_DIR / "missed_doublets_heatmap.png",
        "False negatives by real dataset and method, highlighting where true "
        "doublets were missed.",
    ),
    (
        "Supplementary Figure 5",
        FIG_DIR / "marker_impact_summary.png",
        "Marker-gene stability and effect-size shifts after predicted doublet "
        "removal in pbmc-1A-dm.",
    ),
]


def count_words(text: str) -> int:
    return len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?", text))


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in [
        ("Title", 16, RGBColor(31, 78, 121)),
        ("Heading 1", 12, RGBColor(31, 78, 121)),
        ("Heading 2", 11, RGBColor(31, 78, 121)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def add_meta_block(doc: Document, extra: str | None = None) -> None:
    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(TITLE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"{STUDENT} | {COURSE}").bold = True
    if extra:
        doc.add_paragraph(extra).alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(80, 80, 80)


def load_summary_rows() -> list[dict[str, str]]:
    path = ROOT / "results" / "real" / "real_method_summary.csv"
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def add_results_table(doc: Document) -> None:
    doc.add_heading("Table 1. Real-data benchmark summary", level=2)
    headers = ["Method", "Precision", "Recall", "F1", "AUROC", "AUPRC", "Runtime (s)"]
    rows = load_summary_rows()
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(9)

    for row in rows:
        cells = table.add_row().cells
        values = [
            row["method"],
            f"{float(row['mean_precision']):.3f}",
            f"{float(row['mean_recall']):.3f}",
            f"{float(row['mean_f1']):.3f}",
            f"{float(row['mean_auroc']):.3f}",
            f"{float(row['mean_auprc']):.3f}",
            f"{float(row['mean_runtime_seconds']):.2f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9)
    add_caption(
        doc,
        "Mean metrics across 16 processed real Xi and Li benchmark datasets; "
        "doublets are the positive class.",
    )


def add_picture(doc: Document, heading: str, path: Path, caption: str, width: float = 5.8) -> None:
    doc.add_heading(heading, level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def build_main_report(main_words: int, abstract_words: int) -> Path:
    doc = Document()
    set_doc_defaults(doc)
    add_meta_block(
        doc,
        f"Main text word count: {main_words} | Abstract word count: {abstract_words} | Display items: 2",
    )

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(ABSTRACT)

    for heading, body in MAIN_SECTIONS[:3]:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)

    doc.add_page_break()
    add_results_table(doc)
    composite_path = create_composite_figure()
    add_picture(
        doc,
        "Figure 1. Real-data benchmark and downstream impact",
        composite_path,
        "A, F1 scores by real dataset and method. B, Mean AUROC and AUPRC. C, "
        "Equal-call-budget top-N recovery. D, Representative UMAP before and "
        "after predicted doublet removal.",
        width=6.65,
    )
    doc.add_page_break()

    for heading, body in MAIN_SECTIONS[3:]:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)

    doc.add_heading("References", level=1)
    for ref in REFERENCES:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)

    out = REPORT_DIR / "miniproject7_3127_ica2_main_report.docx"
    doc.save(out)
    return out


def build_supporting_materials(supp_words: int, reflection_words: int) -> Path:
    doc = Document()
    set_doc_defaults(doc)
    add_meta_block(
        doc,
        f"Supporting Methods word count: {supp_words} | Reflection word count: {reflection_words}",
    )
    doc.add_heading("Supporting Methods", level=1)
    doc.add_paragraph(SUPP_METHODS)

    doc.add_heading("Reflection", level=1)
    doc.add_paragraph(REFLECTION)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Supplementary Figures", level=1)
    for label, path, caption in SUPP_FIGURES:
        add_picture(doc, label, path, caption, width=5.8)

    out = REPORT_DIR / "miniproject7_3127_ica2_supporting_materials.docx"
    doc.save(out)
    return out


def write_markdown_copies() -> None:
    main_words = sum(count_words(body) for _, body in MAIN_SECTIONS)
    abstract_words = count_words(ABSTRACT)
    supp_words = count_words(SUPP_METHODS)
    reflection_words = count_words(REFLECTION)

    main_md = [
        f"# {TITLE}",
        "",
        f"{STUDENT} | {COURSE}",
        "",
        f"Main text word count: {main_words}",
        f"Abstract word count: {abstract_words}",
        "Display items: 2",
        "",
        "## Abstract",
        "",
        ABSTRACT,
        "",
    ]
    for heading, body in MAIN_SECTIONS:
        main_md.extend([f"## {heading}", "", body, ""])
    main_md.extend(["## References", ""])
    main_md.extend(f"- {ref}" for ref in REFERENCES)
    (REPORT_DIR / "miniproject7_3127_ica2_main_report.md").write_text(
        "\n".join(main_md), encoding="utf-8"
    )

    supp_md = [
        f"# Supporting Materials: {TITLE}",
        "",
        f"{STUDENT} | {COURSE}",
        "",
        f"Supporting Methods word count: {supp_words}",
        f"Reflection word count: {reflection_words}",
        "",
        "## Supporting Methods",
        "",
        SUPP_METHODS,
        "",
        "## Reflection",
        "",
        REFLECTION,
        "",
        "## Supplementary Figures",
        "",
    ]
    for label, path, caption in SUPP_FIGURES:
        supp_md.extend([f"- {label}: {caption} Source: `{path.relative_to(ROOT)}`"])
    (REPORT_DIR / "miniproject7_3127_ica2_supporting_materials.md").write_text(
        "\n".join(supp_md), encoding="utf-8"
    )


def main() -> None:
    REPORT_DIR.mkdir(exist_ok=True)
    main_words = sum(count_words(body) for _, body in MAIN_SECTIONS)
    abstract_words = count_words(ABSTRACT)
    supp_words = count_words(SUPP_METHODS)
    reflection_words = count_words(REFLECTION)
    checks = [
        ("abstract <=70", abstract_words <= 70, abstract_words),
        ("main text <=1000", main_words <= 1000, main_words),
        ("supporting methods <=500", supp_words <= 500, supp_words),
        ("reflection <=200", reflection_words <= 200, reflection_words),
        ("supp figures <=5", len(SUPP_FIGURES) <= 5, len(SUPP_FIGURES)),
    ]
    for label, ok, value in checks:
        print(f"{label}: {value} ({'OK' if ok else 'FAIL'})")
    if not all(ok for _, ok, _ in checks):
        raise SystemExit("ICA2 length/display constraints failed")

    write_markdown_copies()
    main_docx = build_main_report(main_words, abstract_words)
    supp_docx = build_supporting_materials(supp_words, reflection_words)
    print(f"Wrote {main_docx.relative_to(ROOT)}")
    print(f"Wrote {supp_docx.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
