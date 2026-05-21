"""Build a Word report from the Markdown report draft."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


BODY_FONT = "Arial"
BODY_SIZE = Pt(11)
BODY_LINE_SPACING = 1.15


def configure_paragraph(paragraph, *, size=BODY_SIZE, line_spacing=BODY_LINE_SPACING) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = line_spacing
    paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = BODY_FONT
        run.font.size = size


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = BODY_LINE_SPACING
    normal.paragraph_format.space_after = Pt(6)

    heading_sizes = {
        "Title": 16,
        "Heading 1": 14,
        "Heading 2": 12,
        "Heading 3": 11,
        "Caption": 9,
        "List Bullet": 11,
    }
    for style_name, size in heading_sizes.items():
        style = document.styles[style_name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = BODY_LINE_SPACING
        style.paragraph_format.space_after = Pt(6)


def add_table(document: Document, lines: list[str]) -> None:
    rows = [
        [cell.strip() for cell in line.strip().strip("|").split("|")]
        for line in lines
        if line.strip()
    ]
    rows = [row for row in rows if not all(set(cell) <= {"-", ":"} for cell in row)]
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            paragraph = table.cell(i, j).paragraphs[0]
            run = paragraph.add_run(cell)
            run.font.name = BODY_FONT
            run.font.size = Pt(8)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.space_after = Pt(0)
            if i == 0:
                run.bold = True
    document.add_paragraph()


def add_markdownish(document: Document, text: str) -> None:
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(document, table_lines)
            continue
        else:
            document.add_paragraph(stripped)
        i += 1


def add_figure(document: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    document.add_picture(str(path), width=Inches(5.9))
    paragraph = document.add_paragraph(caption)
    paragraph.style = "Caption"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    configure_paragraph(paragraph, size=Pt(9), line_spacing=1.0)


def split_references(text: str) -> tuple[str, str | None]:
    text = text.replace("\r\n", "\n")
    marker = "\n## References\n"
    if marker not in text:
        return text, None
    body, refs = text.split(marker, 1)
    return body.rstrip(), "## References\n" + refs.strip()


def main() -> None:
    source = Path("report/miniproject7_3127_report.md")
    target = Path("report/miniproject7_3127_report.docx")
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    configure_styles(document)
    body_text, references_text = split_references(source.read_text(encoding="utf-8"))
    add_markdownish(document, body_text)

    figures_heading = document.add_paragraph()
    figures_run = figures_heading.add_run("Figures")
    figures_run.font.name = BODY_FONT
    figures_run.bold = True
    figures_run.font.size = Pt(16)
    configure_paragraph(figures_heading, size=Pt(16), line_spacing=1.15)
    add_figure(
        document,
        Path("figures/pipeline_workflow.png"),
        "Figure 1. Reproducible benchmark workflow.",
    )
    add_figure(
        document,
        Path("figures/f1_precision_recall_comparison.png"),
        "Figure 2. Precision, recall, and F1 comparison across doublet detection methods.",
    )
    add_figure(
        document,
        Path("figures/auprc_auroc_comparison.png"),
        "Figure 3. AUROC and AUPRC comparison across doublet score rankings.",
    )
    add_figure(
        document,
        Path("figures/umap_before_after_doublet_removal.png"),
        "Figure 4. UMAP comparison before and after removing predicted doublets.",
    )
    add_figure(
        document,
        Path("figures/f1_by_dataset_heatmap.png"),
        "Figure 5. F1 scores across simulated benchmark scenarios.",
    )
    add_figure(
        document,
        Path("figures/runtime_comparison.png"),
        "Figure 6. Runtime comparison across methods.",
    )
    add_figure(
        document,
        Path("figures/doublet_type_recall_comparison.png"),
        "Figure 7. Recall stratified by simulated doublet type.",
    )
    add_figure(
        document,
        Path("figures/missed_doublets_heatmap.png"),
        "Figure 8. False negatives by dataset and method.",
    )
    add_figure(
        document,
        Path("figures/real/f1_precision_recall_comparison.png"),
        "Figure 9. Real-data precision, recall, and F1 comparison.",
    )
    add_figure(
        document,
        Path("figures/real/auprc_auroc_comparison.png"),
        "Figure 10. Real-data AUROC and AUPRC comparison.",
    )
    add_figure(
        document,
        Path("figures/real/calibrated_f1_precision_recall_comparison.png"),
        "Figure 11. Real-data calibrated top-N comparison with the same doublet-call budget for every method.",
    )
    add_figure(
        document,
        Path("figures/real/f1_by_dataset_heatmap.png"),
        "Figure 12. Real-data F1 scores by dataset and method.",
    )
    add_figure(
        document,
        Path("figures/real/runtime_comparison.png"),
        "Figure 13. Real-data runtime comparison.",
    )
    add_figure(
        document,
        Path("figures/real/doubletfinder_pk_sensitivity.png"),
        "Figure 14. DoubletFinder pK sensitivity across representative real datasets.",
    )
    add_figure(
        document,
        Path("figures/real/umap_before_after_doublet_removal.png"),
        "Figure 15. Real-data UMAP before and after predicted doublet removal.",
    )
    add_figure(
        document,
        Path("figures/real/missed_doublets_heatmap.png"),
        "Figure 16. Real-data false negatives by dataset and method.",
    )
    add_figure(
        document,
        Path("figures/real/marker_impact_summary.png"),
        "Figure 17. Marker-gene stability and effect-size shifts after predicted doublet removal.",
    )
    if references_text:
        document.add_page_break()
        add_markdownish(document, references_text)

    document.save(target)
    print(f"Wrote {target}")


if __name__ == "__main__":
    main()
